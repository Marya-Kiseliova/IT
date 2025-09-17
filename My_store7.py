# Online_Store_Cards.py

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import pandas as pd
import os
from datetime import datetime
import csv
from docx import Document
import openpyxl
from PIL import Image  # Для фото; установи pip install Pillow

# --- НАСТРОЙКИ ---
DATA_FOLDER = 'StoreData'
if not os.path.exists(DATA_FOLDER):
    os.makedirs(DATA_FOLDER)

INVENTORY_FILE = os.path.join(DATA_FOLDER, 'inventory.csv')
TRANSACTIONS_FILE = os.path.join(DATA_FOLDER, 'transactions.csv')
COSTS_FILE = os.path.join(DATA_FOLDER, 'costs.csv')
CARDS_FOLDER = os.path.join(DATA_FOLDER, 'cards')
if not os.path.exists(CARDS_FOLDER):
    os.makedirs(CARDS_FOLDER)

MARKUP_PERCENTAGE = 30.0  # Наценка для розничной цены

class StoreApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Учет и карточки товаров для интернет-магазина")
        self.root.geometry("900x700")
        self.inventory = self.read_inventory()
        self.setup_files()
        self.create_widgets()

    def setup_files(self):
        if not os.path.exists(INVENTORY_FILE):
            with open(INVENTORY_FILE, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerow(['product_name', 'quantity', 'cost_price', 'sell_price', 'description', 'characteristics', 'photo_url'])

        # ... (остальные файлы как в предыдущем коде)

    def read_inventory(self):
        inventory = {}
        if os.path.exists(INVENTORY_FILE):
            with open(INVENTORY_FILE, 'r', newline='', encoding='utf-8') as f:
                reader = csv.reader(f)
                headers = next(reader)
                for row in reader:
                    if row:
                        inventory[row[0]] = {
                            'quantity': int(float(row[1])),
                            'cost_price': float(row[2]),
                            'sell_price': float(row[3]),
                            'description': row[4] if len(row) > 4 else '',
                            'characteristics': row[5] if len(row) > 5 else '',
                            'photo_url': row[6] if len(row) > 6 else ''
                        }
        return inventory

    def write_inventory(self):
        with open(INVENTORY_FILE, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['product_name', 'quantity', 'cost_price', 'sell_price', 'description', 'characteristics', 'photo_url'])
            for name, data in self.inventory.items():
                writer.writerow([name, data['quantity'], data['cost_price'], data['sell_price'], data['description'], data['characteristics'], data['photo_url']])

    # ... (функции для операций, расходов, отчетов как в предыдущем коде)

    def create_widgets(self):
        style = ttk.Style()
        style.configure("Treeview", font=('Arial', 10), rowheight=25)
        style.configure("Treeview.Heading", font=('Arial', 12, 'bold'))
        style.configure("TLabel", font=('Arial', 11))
        style.configure("TButton", font=('Arial', 11))
        style.configure("TCombobox", font=('Arial', 11))

        notebook = ttk.Notebook(self.root)
        notebook.pack(pady=10, padx=10, fill='both', expand=True)

        # Вкладки для Товаров, Операций, Расходов, Отчетов как в предыдущем

        # Новая вкладка: Карточки товаров
        cards_frame = ttk.LabelFrame(notebook, text="Карточки товаров", padding=10)
        notebook.add(cards_frame, text='Карточки товаров')

        ttk.Label(cards_frame, text="Выберите товар:").grid(row=0, column=0, padx=5, pady=5, sticky='e')
        self.card_product_combo = ttk.Combobox(cards_frame, values=list(self.inventory.keys()), font=('Arial', 11))
        self.card_product_combo.grid(row=0, column=1, padx=5, pady=5)
        self.card_product_combo.bind('<<ComboboxSelected>>', self.load_card_data)

        ttk.Label(cards_frame, text="Описание:").grid(row=1, column=0, padx=5, pady=5, sticky='e')
        self.card_desc_text = tk.Text(cards_frame, height=5, width=40, font=('Arial', 11))
        self.card_desc_text.grid(row=1, column=1, padx=5, pady=5)

        ttk.Label(cards_frame, text="Характеристики (через запятую):").grid(row=2, column=0, padx=5, pady=5, sticky='e')
        self.card_chars_entry = ttk.Entry(cards_frame, font=('Arial', 11))
        self.card_chars_entry.grid(row=2, column=1, padx=5, pady=5)

        ttk.Label(cards_frame, text="Ссылка на фото:").grid(row=3, column=0, padx=5, pady=5, sticky='e')
        self.card_photo_entry = ttk.Entry(cards_frame, font=('Arial', 11))
        self.card_photo_entry.grid(row=3, column=1, padx=5, pady=5)

        ttk.Button(cards_frame, text="Сохранить карточку", command=self.save_card).grid(row=4, column=0, columnspan=2, pady=10)
        ttk.Button(cards_frame, text="Генерировать HTML/PDF", command=self.generate_card).grid(row=5, column=0, columnspan=2, pady=10)

        self.card_tree = ttk.Treeview(cards_frame, columns=('Товар', 'Описание', 'Хар-ки', 'Фото'), show='headings')
        for col in ('Товар', 'Описание', 'Хар-ки', 'Фото'):
            self.card_tree.heading(col, text=col)
            self.card_tree.column(col, width=150, anchor='center')
        self.card_tree.grid(row=6, column=0, columnspan=2, pady=10, sticky='nsew')
        self.update_card_tree()

    def load_card_data(self, event):
        name = self.card_product_combo.get()
        if name in self.inventory:
            data = self.inventory[name]
            self.card_desc_text.delete(1.0, tk.END)
            self.card_desc_text.insert(tk.END, data['description'])
            self.card_chars_entry.delete(0, tk.END)
            self.card_chars_entry.insert(0, data['characteristics'])
            self.card_photo_entry.delete(0, tk.END)
            self.card_photo_entry.insert(0, data['photo_url'])

    def save_card(self):
        name = self.card_product_combo.get()
        if name in self.inventory:
            self.inventory[name]['description'] = self.card_desc_text.get(1.0, tk.END).strip()
            self.inventory[name]['characteristics'] = self.card_chars_entry.get()
            self.inventory[name]['photo_url'] = self.card_photo_entry.get()
            self.write_inventory()
            self.update_card_tree()
            messagebox.showinfo("Успех", f"Карточка '{name}' сохранена.")

    def generate_card(self):
        name = self.card_product_combo.get()
        if name in self.inventory:
            data = self.inventory[name]
            file_path = filedialog.asksaveasfilename(defaultextension=".html", filetypes=[("HTML", "*.html"), ("PDF", "*.pdf")])
            if file_path:
                if file_path.endswith('.html'):
                    self.generate_html_card(name, data, file_path)
                elif file_path.endswith('.pdf'):
                    self.generate_pdf_card(name, data, file_path)
                messagebox.showinfo("Успех", "Карточка сгенерирована.")

    def generate_html_card(self, name, data, file_path):
        html = f"""
        <html>
        <body>
        <h1>{name}</h1>
        <img src="{data['photo_url']}" alt="{name}" width="300">
        <p>Описание: {data['description']}</p>
        <p>Характеристики: {data['characteristics']}</p>
        <p>Цена: {data['sell_price']:.2f} BYN</p>
        </body>
        </html>
        """
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html)

    def generate_pdf_card(self, name, data, file_path):
        doc = Document()
        doc.add_heading(name, 0)
        doc.add_paragraph(f"Описание: {data['description']}")
        doc.add_paragraph(f"Характеристики: {data['characteristics']}")
        doc.add_paragraph(f"Цена: {data['sell_price']:.2f} BYN")
        if data['photo_url'] and os.path.exists(data['photo_url']):
            doc.add_picture(data['photo_url'], width=docx.shared.Inches(4))
        doc.save(file_path)

    def update_card_tree(self):
        for item in self.card_tree.get_children():
            self.card_tree.delete(item)
        for name, data in self.inventory.items():
            self.card_tree.insert('', 'end', values=(name, data['description'][:20] + '...', data['characteristics'][:20] + '...', data['photo_url'][:20] + '...'))

    # ... (остальные функции для учета, операций, расходов, отчетов как в предыдущем коде)

if __name__ == "__main__":
    root = tk.Tk()
    app = StoreApp(root)
    root.mainloop()